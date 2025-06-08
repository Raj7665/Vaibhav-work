from flask import Flask, render_template, request, send_file
from docx import Document
import io

app = Flask(__name__)

def generate_filled_docx(replacements):
    doc = Document("quotation_template.docx")

    def replace_text(paragraphs):
        for p in paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(key, val)

    for para in doc.paragraphs:
        replace_text([para])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell.paragraphs)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@app.route('/', methods=['GET', 'POST'])
def form():
    if request.method == 'POST':
        data = request.form
        replacements = {
            "{{name}}": data.get("name", ""),
            "{{date}}": data.get("date", ""),
            "{{phone_number}}": data.get("phone_number", ""),
            "{{ref_no}}": data.get("ref_no", ""),
            "{{system_title}}": data.get("system_title", ""),
            "{{system_cost}}": data.get("system_cost", ""),
            "{{annual_generation}}": data.get("annual_generation", ""),
            "{{tariff_rate}}": data.get("tariff_rate", ""),
            "{{payback_period}}": data.get("payback_period", "")
            # Add other fields as needed
        }
        docx_file = generate_filled_docx(replacements)
        return send_file(docx_file, as_attachment=True, download_name="filled_quotation.docx")
    return render_template("form.html")
